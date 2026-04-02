#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
ملف واحد لبناء أقوى دمج واسترجاع عملي لمجلد مشروع QTO:
- يقرأ الملفات من المجلد وكل المجلدات الفرعية.
- يحولها إلى مقاطع منظمة مع بيانات وصفية قوية.
- يبني استرجاعًا هجينًا: كلمات + دلالي.
- يدمج النتائج بـ RRF.
- يعيد ترتيبها بـ Cross-Encoder عند توفره.
- يطبق Query Rewrite / Multi-Query.
- يفرض إجابة مقيدة بالأدلة فقط مع امتناع عند ضعف الدليل.
- يحفظ قاعدة معرفة مدمجة داخل نفس المجلد.

الاعتماديات المقترحة:
    pip install sentence-transformers pymupdf openpyxl python-docx numpy

أمثلة تشغيل:
    python qto_ultimate_merge_pipeline.py --folder . --build
    python qto_ultimate_merge_pipeline.py --folder . --ask "ما كمية الخرسانة في القواعد؟"

إذا أردت ربطه مع مزود LLM فعلي، ضع متغيرات البيئة التالية حسب مزودك:
    LLM_BASE_URL
    LLM_API_KEY
    LLM_MODEL

يدعم أي API متوافق مع OpenAI Chat Completions.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import math
import os
import pickle
import re
import statistics
import sys
import textwrap
import time
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

try:
    import numpy as np
except Exception:  # pragma: no cover
    np = None

# اختياري
try:
    from sentence_transformers import CrossEncoder, SentenceTransformer
except Exception:  # pragma: no cover
    CrossEncoder = None
    SentenceTransformer = None

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover
    load_workbook = None

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    import requests
except Exception:  # pragma: no cover
    requests = None


SUPPORTED_EXTENSIONS = {
    ".txt", ".md", ".json", ".jsonl", ".csv", ".tsv",
    ".xlsx", ".xlsm", ".pdf", ".docx"
}

DEFAULT_EMBED_MODEL = "BAAI/bge-m3"
DEFAULT_RERANK_MODEL = "BAAI/bge-reranker-v2-m3"
DEFAULT_INDEX_FILE = ".qto_merged_index.pkl"
DEFAULT_KB_FILE = ".qto_merged_kb.jsonl"
DEFAULT_RESULTS_FILE = ".qto_last_results.json"

QUERY_REWRITE_PROMPT = """You are a retrieval optimizer for construction, BOQ, tendering, quantity surveying, drawings, specifications, and QTO.
Rewrite the user query into 6 search queries that maximize recall without losing precision.
Rules:
1) Preserve quantities, units, trade names, section numbers, drawing IDs, and codes.
2) Generate variants for abbreviations, synonyms, and professional phrasing.
3) Keep each rewrite short and retrieval-friendly.
4) Output valid JSON only.
Schema:
{
  "rewrites": ["...", "..."],
  "filters": {
    "drawing_ids": [],
    "spec_sections": [],
    "trade_keywords": [],
    "unit_keywords": []
  }
}
User query:
{query}
"""

ANSWER_PROMPT = """You are a high-precision QTO evidence engine.
You must answer ONLY from the supplied evidence.
If the evidence is insufficient, contradictory, or low-confidence, refuse to guess and return NEED_REVIEW.

Rules:
1) Never invent quantities, units, item names, section references, or drawing numbers.
2) Prefer table rows and direct quantitative evidence over narrative text.
3) Mention the exact source file and page/sheet/row when available.
4) If two sources conflict, explain the conflict and return NEED_REVIEW unless one source is clearly newer or better matched.
5) Output valid JSON only.

Schema:
{
  "status": "OK" | "NEED_REVIEW",
  "answer": "short grounded answer",
  "confidence": 0.0,
  "evidence": [
    {
      "source": "path-or-file",
      "locator": "page/sheet/row",
      "quote": "short snippet",
      "why_used": "reason"
    }
  ],
  "notes": ["..."]
}

User query:
{query}

Evidence:
{evidence}
"""

LOG_FORMAT = "[%(asctime)s] %(levelname)s - %(message)s"
logging.basicConfig(level=logging.INFO, format=LOG_FORMAT)
logger = logging.getLogger("qto-ultimate-merge")


@dataclass
class Chunk:
    chunk_id: str
    text: str
    source_path: str
    source_name: str
    source_type: str
    page: Optional[int] = None
    sheet: Optional[str] = None
    row: Optional[int] = None
    heading: Optional[str] = None
    trade: Optional[str] = None
    revision: Optional[str] = None
    drawing_id: Optional[str] = None
    unit_candidates: List[str] = field(default_factory=list)
    tokens: List[str] = field(default_factory=list)
    meta: Dict[str, Any] = field(default_factory=dict)


@dataclass
class SearchHit:
    chunk: Chunk
    score: float
    reasons: List[str] = field(default_factory=list)


def sha1_text(text: str) -> str:
    return hashlib.sha1(text.encode("utf-8", errors="ignore")).hexdigest()


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[\t\r]+", " ", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_for_search(text: str) -> str:
    text = text.lower()
    text = text.replace("㎥", " m3 ").replace("㎡", " m2 ")
    text = text.replace("cubic meter", "m3").replace("square meter", "m2")
    text = re.sub(r"[^\w\.\-/#+]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def tokenize(text: str) -> List[str]:
    text = normalize_for_search(text)
    raw = re.findall(r"[a-zA-Z]+\d*|\d+(?:\.\d+)?|[\w\-/#.]+", text)
    tokens = [t for t in raw if len(t) > 1 or t.isdigit()]
    return tokens


def guess_trade(path: Path, text: str) -> Optional[str]:
    hay = f"{str(path).lower()}\n{text[:2500].lower()}"
    trade_map = {
        "concrete": ["concrete", "rc", "blinding", "rebar", "reinforcement", "formwork", "footing", "raft"],
        "architectural": ["architectural", "finish", "blockwork", "plaster", "tiles", "ceiling", "painting"],
        "mechanical": ["mechanical", "hvac", "duct", "pipe", "piping", "pump", "chiller"],
        "electrical": ["electrical", "cable", "lighting", "db", "panel", "earthing", "conduit"],
        "plumbing": ["plumbing", "sanitary", "drainage", "manhole", "water supply"],
        "steel": ["steel", "structural steel", "fabrication", "bolt", "weld"],
    }
    best_trade, best_score = None, 0
    for trade, kws in trade_map.items():
        score = sum(1 for kw in kws if kw in hay)
        if score > best_score:
            best_trade, best_score = trade, score
    return best_trade


def extract_units(text: str) -> List[str]:
    units = re.findall(r"\b(m3|m2|m|kg|ton|tons|nr|no\.?|pcs|item|set|ls|l\.s\.|m\^2|m\^3|mm|cm)\b", text.lower())
    seen = []
    for u in units:
        if u not in seen:
            seen.append(u)
    return seen


def infer_revision(text: str, path: Path) -> Optional[str]:
    hay = f"{path.name} {text[:2000]}"
    patterns = [
        r"\brev(?:ision)?\s*[:\-]?\s*([A-Z0-9_.-]+)",
        r"\bver(?:sion)?\s*[:\-]?\s*([A-Z0-9_.-]+)",
        r"\bissued\s*for\s*([A-Z0-9_. -]+)",
    ]
    for pat in patterns:
        m = re.search(pat, hay, flags=re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return None


def infer_drawing_id(text: str, path: Path) -> Optional[str]:
    hay = f"{path.name}\n{text[:2000]}"
    patterns = [
        r"\b([A-Z]{1,5}-?[A-Z]{0,4}-?\d{2,6}[A-Z0-9-]*)\b",
        r"\b(DWG[- ]?\d{2,6}[A-Z0-9-]*)\b",
        r"\b(SHT[- ]?\d{2,6}[A-Z0-9-]*)\b",
    ]
    for pat in patterns:
        m = re.search(pat, hay)
        if m:
            return m.group(1).strip()
    return None


class LLMClient:
    def __init__(self, base_url: Optional[str] = None, api_key: Optional[str] = None, model: Optional[str] = None):
        self.base_url = base_url or os.getenv("LLM_BASE_URL")
        self.api_key = api_key or os.getenv("LLM_API_KEY")
        self.model = model or os.getenv("LLM_MODEL")

    @property
    def enabled(self) -> bool:
        return bool(self.base_url and self.api_key and self.model and requests)

    def chat_json(self, prompt: str, temperature: float = 0.0, max_tokens: int = 1200) -> Optional[Dict[str, Any]]:
        if not self.enabled:
            return None
        url = self.base_url.rstrip("/") + "/chat/completions"
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        payload = {
            "model": self.model,
            "temperature": temperature,
            "response_format": {"type": "json_object"},
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": max_tokens,
        }
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=120)
            r.raise_for_status()
            data = r.json()
            content = data["choices"][0]["message"]["content"]
            return json.loads(content)
        except Exception as e:
            logger.warning("فشل استدعاء LLM: %s", e)
            return None


class FileLoader:
    def __init__(self, root_folder: str):
        self.root = Path(root_folder).resolve()

    def iter_files(self) -> Iterable[Path]:
        for path in self.root.rglob("*"):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield path

    def load(self) -> List[Chunk]:
        chunks: List[Chunk] = []
        for path in self.iter_files():
            try:
                chunks.extend(self._load_one(path))
            except Exception as e:
                logger.exception("فشل قراءة الملف %s: %s", path, e)
        return chunks

    def _load_one(self, path: Path) -> List[Chunk]:
        ext = path.suffix.lower()
        if ext in {".txt", ".md"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
            return self._chunk_text_document(path, text, source_type=ext[1:])
        if ext in {".json", ".jsonl"}:
            return self._load_json(path)
        if ext in {".csv", ".tsv"}:
            return self._load_csv(path)
        if ext in {".xlsx", ".xlsm"}:
            return self._load_xlsx(path)
        if ext == ".pdf":
            return self._load_pdf(path)
        if ext == ".docx":
            return self._load_docx(path)
        return []

    def _base_meta(self, path: Path, text: str, source_type: str) -> Dict[str, Any]:
        return {
            "source_path": str(path),
            "source_name": path.name,
            "source_type": source_type,
            "trade": guess_trade(path, text),
            "revision": infer_revision(text, path),
            "drawing_id": infer_drawing_id(text, path),
            "unit_candidates": extract_units(text),
        }

    def _make_chunk(
        self,
        path: Path,
        text: str,
        source_type: str,
        page: Optional[int] = None,
        sheet: Optional[str] = None,
        row: Optional[int] = None,
        heading: Optional[str] = None,
        extra_meta: Optional[Dict[str, Any]] = None,
    ) -> Chunk:
        text = normalize_text(text)
        meta = self._base_meta(path, text, source_type)
        if extra_meta:
            meta.update(extra_meta)
        chunk_id = sha1_text(f"{path}|{page}|{sheet}|{row}|{heading}|{text[:1000]}")
        tokens = tokenize(text)
        return Chunk(
            chunk_id=chunk_id,
            text=text,
            source_path=str(path),
            source_name=path.name,
            source_type=source_type,
            page=page,
            sheet=sheet,
            row=row,
            heading=heading,
            trade=meta.get("trade"),
            revision=meta.get("revision"),
            drawing_id=meta.get("drawing_id"),
            unit_candidates=meta.get("unit_candidates", []),
            tokens=tokens,
            meta=meta,
        )

    def _chunk_text_document(self, path: Path, text: str, source_type: str) -> List[Chunk]:
        text = normalize_text(text)
        if not text:
            return []
        lines = text.splitlines()
        sections: List[Tuple[str, List[str]]] = []
        current_heading = "ROOT"
        bucket: List[str] = []

        for line in lines:
            stripped = line.strip()
            if self._looks_like_heading(stripped):
                if bucket:
                    sections.append((current_heading, bucket))
                    bucket = []
                current_heading = stripped[:300]
            else:
                bucket.append(line)
        if bucket:
            sections.append((current_heading, bucket))

        out: List[Chunk] = []
        for heading, body_lines in sections:
            body = normalize_text("\n".join(body_lines))
            if not body:
                continue
            for piece in self._split_long_text(body, max_chars=1800, overlap=250):
                out.append(self._make_chunk(path, piece, source_type=source_type, heading=heading))
        return out

    @staticmethod
    def _looks_like_heading(line: str) -> bool:
        if not line:
            return False
        if len(line) > 120:
            return False
        if re.match(r"^\d+(?:\.\d+)*\s+", line):
            return True
        if re.match(r"^[A-Z][A-Z0-9 /&()\-]{4,}$", line):
            return True
        if re.match(r"^(SECTION|SPEC|DRAWING|BOQ|BILL OF QUANTITIES|GENERAL NOTES)\b", line, flags=re.I):
            return True
        return False

    @staticmethod
    def _split_long_text(text: str, max_chars: int = 1800, overlap: int = 250) -> List[str]:
        if len(text) <= max_chars:
            return [text]
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
        pieces: List[str] = []
        current = ""
        for para in paragraphs:
            if len(current) + len(para) + 2 <= max_chars:
                current = f"{current}\n\n{para}".strip()
            else:
                if current:
                    pieces.append(current)
                if len(para) <= max_chars:
                    current = para
                else:
                    start = 0
                    while start < len(para):
                        end = min(start + max_chars, len(para))
                        pieces.append(para[start:end])
                        if end == len(para):
                            break
                        start = max(end - overlap, start + 1)
                    current = ""
        if current:
            pieces.append(current)
        return pieces

    def _load_json(self, path: Path) -> List[Chunk]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        ext = path.suffix.lower()
        chunks: List[Chunk] = []
        if ext == ".jsonl":
            for idx, line in enumerate(text.splitlines(), start=1):
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except Exception:
                    obj = {"raw": line}
                flat = self._flatten_json(obj)
                chunks.extend(self._chunk_text_document(path, flat, source_type="jsonl"))
        else:
            try:
                obj = json.loads(text)
            except Exception:
                obj = {"raw": text}
            flat = self._flatten_json(obj)
            chunks.extend(self._chunk_text_document(path, flat, source_type="json"))
        return chunks

    def _flatten_json(self, obj: Any, prefix: str = "") -> str:
        lines: List[str] = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                new_prefix = f"{prefix}.{k}" if prefix else str(k)
                lines.append(self._flatten_json(v, new_prefix))
        elif isinstance(obj, list):
            for i, v in enumerate(obj):
                new_prefix = f"{prefix}[{i}]"
                lines.append(self._flatten_json(v, new_prefix))
        else:
            lines.append(f"{prefix}: {obj}")
        return "\n".join(x for x in lines if x)

    def _load_csv(self, path: Path) -> List[Chunk]:
        delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
        chunks: List[Chunk] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
        if not rows:
            return []
        header = rows[0]
        for idx, row in enumerate(rows[1:], start=2):
            row_map = {str(header[i]).strip(): row[i] if i < len(row) else "" for i in range(len(header))}
            text = " | ".join(f"{k}: {v}" for k, v in row_map.items())
            chunks.append(self._make_chunk(path, text, source_type=path.suffix.lower()[1:], row=idx, extra_meta={"row_map": row_map}))
        return chunks

    def _load_xlsx(self, path: Path) -> List[Chunk]:
        if load_workbook is None:
            logger.warning("openpyxl غير متاح، تم تخطي %s", path)
            return []
        wb = load_workbook(filename=path, data_only=True, read_only=True)
        chunks: List[Chunk] = []
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(c).strip() if c is not None else f"col_{i+1}" for i, c in enumerate(rows[0])]
            for idx, row in enumerate(rows[1:], start=2):
                row_map = {header[i]: ("" if i >= len(row) or row[i] is None else str(row[i])) for i in range(len(header))}
                if not any(v for v in row_map.values()):
                    continue
                text = " | ".join(f"{k}: {v}" for k, v in row_map.items())
                chunks.append(self._make_chunk(path, text, source_type="xlsx", sheet=ws.title, row=idx, extra_meta={"row_map": row_map}))
        return chunks

    def _load_pdf(self, path: Path) -> List[Chunk]:
        if fitz is None:
            logger.warning("PyMuPDF غير متاح، تم تخطي %s", path)
            return []
        doc = fitz.open(path)
        chunks: List[Chunk] = []
        for page_idx, page in enumerate(doc, start=1):
            blocks = page.get_text("blocks")
            texts = []
            for block in blocks:
                if len(block) >= 5:
                    txt = normalize_text(str(block[4]))
                    if txt:
                        texts.append(txt)
            page_text = normalize_text("\n\n".join(texts))
            if not page_text:
                continue
            page_chunks = self._chunk_text_document(path, page_text, source_type="pdf")
            for c in page_chunks:
                c.page = page_idx
            chunks.extend(page_chunks)
        return chunks

    def _load_docx(self, path: Path) -> List[Chunk]:
        if DocxDocument is None:
            logger.warning("python-docx غير متاح، تم تخطي %s", path)
            return []
        doc = DocxDocument(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paras)
        chunks = self._chunk_text_document(path, text, source_type="docx")
        # استخراج الجداول كسطور منفصلة لرفع الدقة
        for table_idx, table in enumerate(doc.tables, start=1):
            rows = []
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [normalize_text(cell.text) for cell in row.cells]
                if any(cells):
                    rows.append((row_idx, cells))
            if rows:
                header = rows[0][1]
                for row_idx, cells in rows[1:]:
                    row_map = {header[i] if i < len(header) and header[i] else f"col_{i+1}": cells[i] if i < len(cells) else "" for i in range(max(len(header), len(cells)))}
                    text = " | ".join(f"{k}: {v}" for k, v in row_map.items())
                    chunks.append(self._make_chunk(path, text, source_type="docx", row=row_idx, extra_meta={"table_index": table_idx, "row_map": row_map}))
        return chunks


class BM25Index:
    def __init__(self, chunks: Sequence[Chunk], k1: float = 1.5, b: float = 0.75):
        self.chunks = list(chunks)
        self.k1 = k1
        self.b = b
        self.doc_len = [len(c.tokens) or 1 for c in self.chunks]
        self.avgdl = statistics.mean(self.doc_len) if self.doc_len else 1.0
        self.df = Counter()
        self.tf: List[Counter] = []
        for c in self.chunks:
            freqs = Counter(c.tokens)
            self.tf.append(freqs)
            for t in freqs.keys():
                self.df[t] += 1
        self.N = len(self.chunks)

    def idf(self, term: str) -> float:
        n_q = self.df.get(term, 0)
        return math.log(1 + ((self.N - n_q + 0.5) / (n_q + 0.5)))

    def search(self, query: str, top_k: int = 30, filters: Optional[Dict[str, Any]] = None) -> List[SearchHit]:
        q_tokens = tokenize(query)
        scores: List[Tuple[int, float, List[str]]] = []
        for idx, chunk in enumerate(self.chunks):
            if not metadata_match(chunk, filters):
                continue
            score = 0.0
            reasons = []
            freqs = self.tf[idx]
            dl = self.doc_len[idx]
            for t in q_tokens:
                if t not in freqs:
                    continue
                idf = self.idf(t)
                tf = freqs[t]
                part = idf * ((tf * (self.k1 + 1)) / (tf + self.k1 * (1 - self.b + self.b * dl / self.avgdl)))
                score += part
                reasons.append(f"bm25:{t}")
            if score > 0:
                scores.append((idx, score, reasons))
        scores.sort(key=lambda x: x[1], reverse=True)
        return [SearchHit(chunk=self.chunks[i], score=s, reasons=r) for i, s, r in scores[:top_k]]


class DenseIndex:
    def __init__(self, chunks: Sequence[Chunk], model_name: str = DEFAULT_EMBED_MODEL):
        self.chunks = list(chunks)
        self.model_name = model_name
        self.model = None
        self.embeddings = None
        if SentenceTransformer is not None and np is not None:
            try:
                self.model = SentenceTransformer(model_name)
                texts = [build_dense_text(c) for c in self.chunks]
                self.embeddings = self.model.encode(texts, normalize_embeddings=True, show_progress_bar=True)
            except Exception as e:
                logger.warning("تعذر بناء DenseIndex: %s", e)
                self.model = None
                self.embeddings = None

    @property
    def enabled(self) -> bool:
        return self.model is not None and self.embeddings is not None and np is not None

    def search(self, query: str, top_k: int = 30, filters: Optional[Dict[str, Any]] = None) -> List[SearchHit]:
        if not self.enabled:
            return []
        q = self.model.encode([query], normalize_embeddings=True)[0]
        sims = np.dot(self.embeddings, q)
        results: List[Tuple[int, float]] = []
        for idx, sim in enumerate(sims):
            chunk = self.chunks[idx]
            if not metadata_match(chunk, filters):
                continue
            results.append((idx, float(sim)))
        results.sort(key=lambda x: x[1], reverse=True)
        out = []
        for idx, score in results[:top_k]:
            out.append(SearchHit(chunk=self.chunks[idx], score=score, reasons=["dense"]))
        return out


class Reranker:
    def __init__(self, model_name: str = DEFAULT_RERANK_MODEL):
        self.model = None
        if CrossEncoder is not None:
            try:
                self.model = CrossEncoder(model_name)
            except Exception as e:
                logger.warning("تعذر تحميل reranker: %s", e)
                self.model = None

    @property
    def enabled(self) -> bool:
        return self.model is not None

    def rerank(self, query: str, hits: List[SearchHit], top_k: int = 12) -> List[SearchHit]:
        if not hits:
            return []
        if not self.enabled:
            return hits[:top_k]
        pairs = [(query, build_rerank_text(h.chunk)) for h in hits]
        try:
            scores = self.model.predict(pairs)
            for h, s in zip(hits, scores):
                h.score = float(s)
                h.reasons.append("rerank")
            hits.sort(key=lambda x: x.score, reverse=True)
        except Exception as e:
            logger.warning("فشل rerank: %s", e)
        return hits[:top_k]


class UltimateQTOEngine:
    def __init__(
        self,
        folder: str,
        embed_model: str = DEFAULT_EMBED_MODEL,
        rerank_model: str = DEFAULT_RERANK_MODEL,
    ):
        self.folder = Path(folder).resolve()
        self.embed_model = embed_model
        self.rerank_model = rerank_model
        self.index_path = self.folder / DEFAULT_INDEX_FILE
        self.kb_path = self.folder / DEFAULT_KB_FILE
        self.results_path = self.folder / DEFAULT_RESULTS_FILE
        self.chunks: List[Chunk] = []
        self.bm25: Optional[BM25Index] = None
        self.dense: Optional[DenseIndex] = None
        self.reranker: Optional[Reranker] = None
        self.llm = LLMClient()

    def build(self, force: bool = False) -> None:
        if self.index_path.exists() and not force:
            self.load_index()
            return
        loader = FileLoader(str(self.folder))
        chunks = loader.load()
        chunks = self._dedupe_and_enrich(chunks)
        self.chunks = chunks
        self.bm25 = BM25Index(self.chunks)
        self.dense = DenseIndex(self.chunks, model_name=self.embed_model)
        self.reranker = Reranker(self.rerank_model)
        self._save_index()
        self._save_jsonl_kb()
        logger.info("تم بناء الفهرس: %s مقطع", len(self.chunks))

    def load_index(self) -> None:
        with self.index_path.open("rb") as f:
            payload = pickle.load(f)
        self.chunks = payload["chunks"]
        self.bm25 = BM25Index(self.chunks)
        self.dense = DenseIndex(self.chunks, model_name=payload.get("embed_model", self.embed_model))
        self.reranker = Reranker(payload.get("rerank_model", self.rerank_model))
        logger.info("تم تحميل الفهرس من %s", self.index_path)

    def ask(self, query: str, top_k: int = 12) -> Dict[str, Any]:
        if not self.chunks:
            self.build(force=False)
        rewrites, filters = self.rewrite_query(query)
        raw_hits = self.retrieve(query, rewrites, filters)
        reranked = self.rerank(query, raw_hits, top_k=top_k)
        validated = self.apply_validation(query, reranked)
        answer = self.answer(query, validated)
        self.results_path.write_text(json.dumps(answer, ensure_ascii=False, indent=2), encoding="utf-8")
        return answer

    def rewrite_query(self, query: str) -> Tuple[List[str], Dict[str, Any]]:
        base_rewrites = heuristic_rewrites(query)
        filters: Dict[str, Any] = {}
        data = self.llm.chat_json(QUERY_REWRITE_PROMPT.format(query=query), temperature=0.0) if self.llm.enabled else None
        if isinstance(data, dict):
            llm_rewrites = [x.strip() for x in data.get("rewrites", []) if isinstance(x, str) and x.strip()]
            base_rewrites.extend(llm_rewrites)
            if isinstance(data.get("filters"), dict):
                filters = data["filters"]
        all_queries = dedupe_preserve([query] + base_rewrites)
        return all_queries[:8], filters

    def retrieve(self, query: str, rewrites: List[str], filters: Optional[Dict[str, Any]] = None) -> List[SearchHit]:
        if self.bm25 is None:
            raise RuntimeError("BM25 غير جاهز")
        dense_hits_all: List[List[SearchHit]] = []
        bm25_hits_all: List[List[SearchHit]] = []

        for q in rewrites:
            bm25_hits_all.append(self.bm25.search(q, top_k=40, filters=filters))
            if self.dense is not None and self.dense.enabled:
                dense_hits_all.append(self.dense.search(q, top_k=40, filters=filters))

        merged = rrf_fuse(
            result_sets=[*bm25_hits_all, *dense_hits_all],
            k=60,
            meta_boost_fn=lambda ch: metadata_boost(query, ch),
        )
        return diversify_hits(merged, max_per_source=4)

    def rerank(self, query: str, hits: List[SearchHit], top_k: int = 12) -> List[SearchHit]:
        if self.reranker is None:
            self.reranker = Reranker(self.rerank_model)
        return self.reranker.rerank(query, hits, top_k=top_k)

    def apply_validation(self, query: str, hits: List[SearchHit]) -> List[SearchHit]:
        query_units = set(extract_units(query))
        query_tokens = set(tokenize(query))
        final: List[SearchHit] = []
        for h in hits:
            penalty = 0.0
            reasons = []
            if query_units and h.chunk.unit_candidates:
                if not query_units.intersection(set(h.chunk.unit_candidates)):
                    penalty += 0.15
                    reasons.append("unit_mismatch")
            if h.chunk.revision:
                reasons.append(f"revision:{h.chunk.revision}")
            lexical_overlap = len(query_tokens.intersection(set(h.chunk.tokens)))
            if lexical_overlap == 0:
                penalty += 0.08
                reasons.append("weak_lexical_overlap")
            if len(h.chunk.text) < 40:
                penalty += 0.05
                reasons.append("too_short")
            h.score = h.score - penalty
            h.reasons.extend(reasons)
            final.append(h)
        final.sort(key=lambda x: x.score, reverse=True)
        return final

    def answer(self, query: str, hits: List[SearchHit]) -> Dict[str, Any]:
        evidence = build_evidence_payload(hits)
        confidence = compute_confidence(hits)

        if confidence < 0.42 or not hits:
            return {
                "status": "NEED_REVIEW",
                "answer": "الأدلة المسترجعة غير كافية أو ليست حاسمة بما يكفي لإخراج إجابة نهائية.",
                "confidence": round(confidence, 4),
                "evidence": evidence,
                "notes": [
                    "تم تفعيل الامتناع عن التخمين.",
                    "راجع أفضل المقاطع المسترجعة أو وسّع البيانات الوصفية أو حسّن بنية القراءة."
                ]
            }

        data = self.llm.chat_json(
            ANSWER_PROMPT.format(query=query, evidence=json.dumps(evidence, ensure_ascii=False, indent=2)),
            temperature=0.0,
            max_tokens=1400,
        ) if self.llm.enabled else None

        if isinstance(data, dict) and data.get("status") in {"OK", "NEED_REVIEW"}:
            data["confidence"] = round(max(float(data.get("confidence", confidence)), confidence), 4)
            data.setdefault("evidence", evidence)
            return data

        # بديل حتمي بدون LLM
        best = hits[0]
        locator = locate(best.chunk)
        return {
            "status": "OK" if confidence >= 0.55 else "NEED_REVIEW",
            "answer": fallback_grounded_answer(query, hits),
            "confidence": round(confidence, 4),
            "evidence": evidence,
            "notes": [
                f"أفضل شاهد من {best.chunk.source_name} عند {locator}",
                "هذه إجابة مقيدة بالأدلة المسترجعة فقط."
            ]
        }

    def _dedupe_and_enrich(self, chunks: List[Chunk]) -> List[Chunk]:
        seen = set()
        out = []
        for c in chunks:
            key = (normalize_for_search(c.text), c.source_name, c.page, c.sheet, c.row)
            if key in seen:
                continue
            seen.add(key)
            if not c.trade:
                c.trade = guess_trade(Path(c.source_path), c.text)
            if not c.unit_candidates:
                c.unit_candidates = extract_units(c.text)
            if not c.tokens:
                c.tokens = tokenize(c.text)
            out.append(c)
        return out

    def _save_index(self) -> None:
        payload = {
            "chunks": self.chunks,
            "embed_model": self.embed_model,
            "rerank_model": self.rerank_model,
            "built_at": time.time(),
        }
        with self.index_path.open("wb") as f:
            pickle.dump(payload, f)

    def _save_jsonl_kb(self) -> None:
        with self.kb_path.open("w", encoding="utf-8") as f:
            for c in self.chunks:
                f.write(json.dumps(asdict(c), ensure_ascii=False) + "\n")


def metadata_match(chunk: Chunk, filters: Optional[Dict[str, Any]]) -> bool:
    if not filters:
        return True
    drawing_ids = {x.lower() for x in filters.get("drawing_ids", []) if isinstance(x, str)}
    spec_sections = {x.lower() for x in filters.get("spec_sections", []) if isinstance(x, str)}
    trade_keywords = {x.lower() for x in filters.get("trade_keywords", []) if isinstance(x, str)}
    unit_keywords = {x.lower() for x in filters.get("unit_keywords", []) if isinstance(x, str)}

    if drawing_ids and (chunk.drawing_id or "").lower() not in drawing_ids:
        return False
    if trade_keywords and (chunk.trade or "").lower() not in trade_keywords:
        return False
    if unit_keywords and not unit_keywords.intersection(set(u.lower() for u in chunk.unit_candidates)):
        return False
    if spec_sections:
        hay = normalize_for_search(chunk.text[:500])
        if not any(sec in hay for sec in spec_sections):
            return False
    return True


def build_dense_text(chunk: Chunk) -> str:
    meta = [
        f"source={chunk.source_name}",
        f"trade={chunk.trade or ''}",
        f"drawing={chunk.drawing_id or ''}",
        f"heading={chunk.heading or ''}",
        f"page={chunk.page or ''}",
        f"sheet={chunk.sheet or ''}",
    ]
    return "\n".join(meta) + "\n" + chunk.text


def build_rerank_text(chunk: Chunk) -> str:
    loc = locate(chunk)
    return f"{chunk.source_name} | {loc} | {chunk.heading or ''}\n{chunk.text}"


def rrf_fuse(
    result_sets: List[List[SearchHit]],
    k: int = 60,
    meta_boost_fn=None,
) -> List[SearchHit]:
    merged_scores: Dict[str, float] = defaultdict(float)
    representative: Dict[str, SearchHit] = {}
    reasons_map: Dict[str, List[str]] = defaultdict(list)

    for result_set in result_sets:
        for rank, hit in enumerate(result_set, start=1):
            cid = hit.chunk.chunk_id
            merged_scores[cid] += 1.0 / (k + rank)
            representative[cid] = hit
            reasons_map[cid].extend(hit.reasons)

    hits: List[SearchHit] = []
    for cid, score in merged_scores.items():
        hit = representative[cid]
        if meta_boost_fn:
            score += float(meta_boost_fn(hit.chunk))
        hit.score = score
        hit.reasons = dedupe_preserve(reasons_map[cid])
        hits.append(hit)

    hits.sort(key=lambda x: x.score, reverse=True)
    return hits


def metadata_boost(query: str, chunk: Chunk) -> float:
    boost = 0.0
    q = normalize_for_search(query)
    if chunk.trade and chunk.trade.lower() in q:
        boost += 0.03
    if chunk.drawing_id and chunk.drawing_id.lower() in q:
        boost += 0.08
    if chunk.unit_candidates and any(u in q for u in chunk.unit_candidates):
        boost += 0.02
    if chunk.source_type in {"xlsx", "csv", "tsv"}:
        boost += 0.025
    if chunk.page is not None:
        boost += 0.005
    return boost


def diversify_hits(hits: List[SearchHit], max_per_source: int = 4) -> List[SearchHit]:
    counts = Counter()
    out = []
    for h in hits:
        key = h.chunk.source_name
        if counts[key] >= max_per_source:
            continue
        counts[key] += 1
        out.append(h)
    return out


def dedupe_preserve(items: Sequence[str]) -> List[str]:
    seen = set()
    out = []
    for item in items:
        if item not in seen:
            seen.add(item)
            out.append(item)
    return out


def heuristic_rewrites(query: str) -> List[str]:
    q = normalize_text(query)
    rewrites = [q]
    replacements = [
        (r"\bqty\b", "quantity"),
        (r"\bboq\b", "bill of quantities"),
        (r"\bconc\b", "concrete"),
        (r"\brc\b", "reinforced concrete"),
        (r"\brebar\b", "reinforcement"),
        (r"\bmep\b", "mechanical electrical plumbing"),
    ]
    q2 = q.lower()
    for pat, repl in replacements:
        q2 = re.sub(pat, repl, q2, flags=re.I)
    rewrites.append(q2)

    units = extract_units(q)
    if units:
        for u in units:
            if u == "m3":
                rewrites.append(q.replace("m3", "cubic meter"))
            elif u == "m2":
                rewrites.append(q.replace("m2", "square meter"))

    tokens = tokenize(q)
    if tokens:
        important = " ".join(tokens[:8])
        rewrites.append(important)
    return dedupe_preserve([x.strip() for x in rewrites if x.strip()])


def locate(chunk: Chunk) -> str:
    parts = []
    if chunk.page is not None:
        parts.append(f"page {chunk.page}")
    if chunk.sheet:
        parts.append(f"sheet {chunk.sheet}")
    if chunk.row is not None:
        parts.append(f"row {chunk.row}")
    if chunk.heading:
        parts.append(f"heading {chunk.heading[:80]}")
    return " | ".join(parts) if parts else "source"


def build_evidence_payload(hits: List[SearchHit], max_items: int = 8) -> List[Dict[str, Any]]:
    payload = []
    for h in hits[:max_items]:
        payload.append({
            "source": h.chunk.source_path,
            "locator": locate(h.chunk),
            "quote": h.chunk.text[:500],
            "why_used": ", ".join(dedupe_preserve(h.reasons))[:300],
            "score": round(h.score, 6),
            "drawing_id": h.chunk.drawing_id,
            "trade": h.chunk.trade,
            "revision": h.chunk.revision,
            "units": h.chunk.unit_candidates,
        })
    return payload


def compute_confidence(hits: List[SearchHit]) -> float:
    if not hits:
        return 0.0
    top_scores = [max(0.0, h.score) for h in hits[:5]]
    if not top_scores:
        return 0.0
    max_score = max(top_scores)
    if max_score <= 0:
        return 0.0
    normalized = [min(1.0, s / max_score) for s in top_scores]
    source_diversity = len({h.chunk.source_name for h in hits[:5]}) / min(5, len(hits[:5]))
    avg_top = sum(normalized) / len(normalized)
    confidence = 0.7 * avg_top + 0.3 * source_diversity
    return round(min(0.99, confidence), 4)


def fallback_grounded_answer(query: str, hits: List[SearchHit]) -> str:
    top = hits[:3]
    snippets = []
    for h in top:
        loc = locate(h.chunk)
        snippets.append(f"{h.chunk.source_name} ({loc}): {truncate(h.chunk.text, 180)}")
    return "أفضل الأدلة المسترجعة للسؤال هي: " + " || ".join(snippets)


def truncate(text: str, n: int) -> str:
    text = normalize_text(text)
    return text if len(text) <= n else text[: n - 3] + "..."


def print_json(data: Dict[str, Any]) -> None:
    print(json.dumps(data, ensure_ascii=False, indent=2))


def main() -> int:
    parser = argparse.ArgumentParser(description="محرك دمج واسترجاع عالي الدقة لمشاريع QTO")
    parser.add_argument("--folder", required=True, help="مسار مجلد المشروع")
    parser.add_argument("--build", action="store_true", help="بناء الفهرس المدمج")
    parser.add_argument("--force", action="store_true", help="إعادة البناء حتى لو الفهرس موجود")
    parser.add_argument("--ask", type=str, help="سؤال على قاعدة المعرفة المدمجة")
    parser.add_argument("--embed-model", default=DEFAULT_EMBED_MODEL)
    parser.add_argument("--rerank-model", default=DEFAULT_RERANK_MODEL)
    args = parser.parse_args()

    engine = UltimateQTOEngine(
        folder=args.folder,
        embed_model=args.embed_model,
        rerank_model=args.rerank_model,
    )

    if args.build or not engine.index_path.exists():
        engine.build(force=args.force)

    if args.ask:
        result = engine.ask(args.ask)
        print_json(result)
        return 0

    info = {
        "status": "OK",
        "message": "تم تجهيز الملف. استخدم --build للبناء و --ask للاستعلام.",
        "index_file": str(engine.index_path),
        "kb_file": str(engine.kb_path),
        "results_file": str(engine.results_path),
    }
    print_json(info)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
